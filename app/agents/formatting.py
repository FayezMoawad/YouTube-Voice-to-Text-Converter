from pathlib import Path
from docx import Document
from app.core.schemas import Transcript, Segment
from app.utils.config import OUTPUT_DIR
import logging

logger = logging.getLogger(__name__)

class FormattingAgent:
    def format_docx(self, transcript: Transcript, include_timestamps: bool = True, include_speakers: bool = True) -> Path:
        doc = Document()
        doc.add_heading(transcript.metadata.title, 0)

        doc.add_paragraph(f"Source: {transcript.metadata.url}")
        doc.add_paragraph(f"Duration: {transcript.metadata.duration} seconds")
        doc.add_paragraph(f"Date: {transcript.metadata.upload_date or 'N/A'}")
        
        doc.add_heading('Transcript', level=1)
        
        for segment in transcript.segments:
            p = doc.add_paragraph()
            
            # Construct prefix: [10s] Speaker: 
            prefix = ""
            if include_timestamps:
                prefix += f"[{segment.start_time:.1f}s] "
            
            if include_speakers:
                prefix += f"{segment.speaker}: "
            
            if prefix:
                runner = p.add_run(prefix)
                runner.bold = True
            
            p.add_run(segment.text)
            
        filename = f"{transcript.metadata.title[:50]}_transcript.docx"
        # Sanitize filename
        filename = "".join([c for c in filename if c.isalpha() or c.isdigit() or c==' ' or c=='_']).rstrip() + ".docx"
        output_path = OUTPUT_DIR / filename
        
        doc.save(str(output_path))
        logger.info(f"DOCX saved to {output_path}")
        return output_path

    def format_txt(self, transcript: Transcript, include_timestamps: bool = True, include_speakers: bool = True) -> Path:
        filename = f"{transcript.metadata.title[:50]}_transcript.txt"
        filename = "".join([c for c in filename if c.isalpha() or c.isdigit() or c==' ' or c=='_']).rstrip() + ".txt"
        output_path = OUTPUT_DIR / filename
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"Title: {transcript.metadata.title}\n")
            f.write(f"URL: {transcript.metadata.url}\n\n")
            for segment in transcript.segments:
                line = ""
                if include_timestamps:
                    line += f"[{segment.start_time:.1f}s] "
                if include_speakers:
                    line += f"{segment.speaker}: "
                
                line += segment.text
                f.write(f"{line}\n")
                
        logger.info(f"TXT saved to {output_path}")
        return output_path
